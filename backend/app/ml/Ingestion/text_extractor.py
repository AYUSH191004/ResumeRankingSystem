"""
Text extraction from PDF and DOCX files
"""
import os
from typing import Optional
import PyPDF2
from docx import Document
import pdfplumber
import logging

logger = logging.getLogger(__name__)


class ResumeTextExtractor:
    """Extract text from resume files"""
    
    def __init__(self):
        self.supported_formats = ['.pdf', '.docx']
    
    def extract(self, file_path: str) -> Optional[str]:
        """
        Extract text from file
        
        Args:
            file_path: Path to the resume file
            
        Returns:
            Extracted text or None if failed
        """
        if not os.path.exists(file_path):
            logger.error(f"File not found: {file_path}")
            return None
        
        file_ext = os.path.splitext(file_path)[1].lower()
        
        try:
            if file_ext == '.pdf':
                return self._extract_from_pdf(file_path)
            elif file_ext == '.docx':
                return self._extract_from_docx(file_path)
            else:
                logger.error(f"Unsupported format: {file_ext}")
                return None
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {str(e)}")
            return None
    
    def _extract_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF using multiple methods"""
        text = ""
        
        # Method 1: Try pdfplumber (better for complex layouts)
        try:
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            
            if text.strip():
                logger.info(f"Extracted {len(text)} chars using pdfplumber")
                return text
        except Exception as e:
            logger.warning(f"pdfplumber failed: {str(e)}")
        
        # Method 2: Fallback to PyPDF2
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            
            logger.info(f"Extracted {len(text)} chars using PyPDF2")
            return text
        except Exception as e:
            logger.error(f"PyPDF2 also failed: {str(e)}")
            return ""
    
    def _extract_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX"""
        try:
            doc = Document(file_path)
            
            # Extract from paragraphs
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            
            # Extract from tables
            tables_text = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text.strip() for cell in row.cells])
                    if row_text.strip():
                        tables_text.append(row_text)
            
            text = "\n".join(paragraphs + tables_text)
            logger.info(f"Extracted {len(text)} chars from DOCX")
            return text
        
        except Exception as e:
            logger.error(f"Error extracting DOCX: {str(e)}")
            return ""
    
    def is_valid_resume(self, text: str, min_length: int = 100) -> bool:
        """
        Check if extracted text looks like a valid resume
        
        Args:
            text: Extracted text
            min_length: Minimum character count
            
        Returns:
            True if valid resume
        """
        if not text or len(text.strip()) < min_length:
            return False
        
        # Check for common resume keywords
        resume_indicators = [
            'experience', 'education', 'skills', 'work',
            'university', 'college', 'email', 'phone',
            'project', 'internship', 'certificate'
        ]
        
        text_lower = text.lower()
        found_indicators = sum(1 for keyword in resume_indicators if keyword in text_lower)
        
        return found_indicators >= 2  # At least 2 indicators
