import os
import re
import logging
from typing import Dict, Optional, List
from PyPDF2 import PdfReader
from docx import Document
from sqlalchemy.orm import declarative_base
from sqlalchemy import create_engine, Column, Integer, String, Float, ForeignKey, Text, Enum, Date, DateTime,JSON
from sqlalchemy.orm import sessionmaker
import datetime 
from transformers import pipeline
import spacy
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import numpy as np

# Common resume section patterns
SECTION_PATTERNS = {
    'experience': r'(work|employment|experience) history?',
    'education': r'educations?|degrees?|academics?',
    'skills': r'(technical|key)?\s*skills?',
    'projects': r'projects?|assignments?'
}

# Enhanced skill keywords list
SKILL_KEYWORDS = [
    # Programming
    'python', 'java', 'c++', 'c#', 'javascript', 'typescript',
    'ruby', 'php', 'swift', 'kotlin', 'go', 'rust',
    
    # Web
    'html', 'css', 'react', 'angular', 'vue', 'django', 'flask',
    'node.js', 'express', 'spring', 'laravel',
    
    # Databases
    'sql', 'mysql', 'postgresql', 'mongodb', 'oracle', 'redis',
    
    # Data Science
    'machine learning', 'deep learning', 'tensorflow', 'pytorch',
    'data analysis', 'pandas', 'numpy', 'spark', 'hadoop',
    
    # DevOps
    'docker', 'kubernetes', 'aws', 'azure', 'gcp', 'ci/cd',
    'terraform', 'ansible', 'jenkins',
    
    # Other
    'project management', 'agile', 'scrum', 'git', 'linux'
]

# Contact info patterns
CONTACT_PATTERNS = {
    'email': r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b',
    'phone': r'(\+\d{1,3}[-\.\s]?)?(\(\d{3}\)|\d{3})[-\.\s]?\d{3}[-\.\s]?\d{4}'
}

# Date patterns for education/experience
DATE_PATTERNS = {
    'year_range': r'(20\d{2}\s*[-–]\s*20\d{2}|present|current)',
    'month_year': r'(jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z]*\s+20\d{2}'
}


# Database model
Base = declarative_base()

class Resume(Base):
    __tablename__ = 'resumes'
    
    id = Column(Integer, primary_key=True)
    file_name = Column(String(255))
    raw_text = Column(Text)
    extracted_data = Column(JSON)
    processed_at = Column(DateTime, default=datetime.datetime.now)
    
    def __repr__(self):
        return f"<Resume(file_name='{self.file_name}')>"

class DocumentProcessor:
    """A class to process PDF/DOCX files, extract text, analyze with NLP and store in DB."""
    
    def __init__(self, input_directory: str = None, output_directory: str = None, db_uri: str = None):
        self.logger = logging.getLogger(__name__)
        self.input_directory = input_directory
        self.output_directory = output_directory
        
        # Initialize database connection
        if db_uri:
            self.engine = create_engine(db_uri)
            Base.metadata.create_all(self.engine)
            self.Session = sessionmaker(bind=self.engine)
        
        if output_directory and not os.path.exists(output_directory):
            os.makedirs(output_directory)

        # Define ORM models for the schema tables
        from sqlalchemy import Table, MetaData, select
        metadata = MetaData()
        metadata.reflect(bind=self.engine)

        self.candidates_table = metadata.tables.get('candidates')
        self.skills_table = metadata.tables.get('skills')
        self.projects_table = metadata.tables.get('projects')
        self.work_experience_table = metadata.tables.get('work_experience')
        self.certifications_table = metadata.tables.get('certifications')
        self.rankings_table = metadata.tables.get('rankings')
        self.analysis_results_table = metadata.tables.get('analysis_results')

    def process_pdf(self, file_path: str) -> Optional[dict]:
        """Enhanced PDF extraction with section detection"""
        try:
            text = ""
            with open(file_path, 'rb') as file:
                reader = PdfReader(file)
                for page in reader.pages:
                    text += self._clean_text(page.extract_text()) + "\n"
            return self._analyze_resume(text)
        except Exception as e:
            self.logger.error(f"PDF processing error: {str(e)}")
            return None
            
    def process_docx(self, file_path: str) -> Optional[dict]:
        """Enhanced DOCX extraction with section detection"""
        try:
            doc = Document(file_path)
            text = "\n".join([self._clean_text(para.text) for para in doc.paragraphs])
            return self._analyze_resume(text)
        except Exception as e:
            self.logger.error(f"DOCX processing error: {str(e)}")
            return None
            
    def _clean_text(self, text: str) -> str:
        """Clean and normalize text"""
        text = re.sub(r'\s+', ' ', text)  # Remove extra whitespace
        text = re.sub(r'[^\x00-\x7F]+', ' ', text)  # Remove non-ASCII
        return text.strip()

    def _analyze_with_nlp(self, text: str) -> dict:
        """Perform NLP analysis on resume text"""
        model=spacy.load('en_core_web_lg')
        doc = model(text)
        
        # Extract entities
        entities = {
            'PERSON': [],
            'ORG': [],
            'DATE': [],
            'CGPA': [],
            'TECH': []
        }
        
        for ent in doc.ents:
            if ent.label_ in entities:
                entities[ent.label_].append(ent.text)
        
        # Extract skills using noun chunks and matching
        skills = []
        for chunk in doc.noun_chunks:
            if any(skill.lower() in chunk.text.lower() for skill in SKILL_KEYWORDS):
                skills.append(chunk.text)
        
        return {
            'entities': entities,
            'skills': list(set(skills))
        }

    def _analyze_resume(self, text: str) -> dict:
        """Enhanced resume analysis with NLP, contact info and dates"""
        # Detect sections
        sections = {
            name: bool(re.search(pattern, text, re.I))
            for name, pattern in SECTION_PATTERNS.items()
        }
        
        # Extract skills
        skills = [
            skill for skill in SKILL_KEYWORDS
            if re.search(rf'\b{re.escape(skill)}\b', text, re.I)
        ]
        
        # Extract contact info
        contacts = {
            'emails': list(set(re.findall(CONTACT_PATTERNS['email'], text, re.I))),
            'phones': list(set(re.findall(CONTACT_PATTERNS['phone'], text)))
        }
        
        # Find dates in experience/education sections
        dates = {
            'education': re.findall(DATE_PATTERNS['year_range'], text, re.I),
            'experience': re.findall(DATE_PATTERNS['month_year'], text, re.I)
        }
        
        # Basic analysis
        analysis = {
            'raw_text': text,
            'sections': sections,
            'skills': skills,
            'contacts': contacts,
            'dates': dates
        }
        
        # Add NLP analysis
        nlp_analysis = self._analyze_with_nlp(text)
        analysis.update(nlp_analysis)
        
        return analysis

    def extract_from_pdf(self, file_path: str) -> str:
        """Legacy method - use process_pdf instead"""
        result = self.process_pdf(file_path)
        return result['raw_text'] if result else ""

    def extract_from_docx(self, file_path: str) -> str:
        """Legacy method - use process_docx instead"""
        result = self.process_docx(file_path)
        return result['raw_text'] if result else ""

    def save_to_db(self, file_name: str, analysis_data: dict):
        """Save analysis results to database"""
        if not hasattr(self, 'Session'):
            self.logger.warning("No database connection configured")
            return False
            
        session = self.Session()
        try:
            resume = Resume(
                file_name=file_name,
                raw_text=analysis_data['raw_text'],
                extracted_data=analysis_data
            )
            session.add(resume)
            session.commit()
            return True
        except Exception as e:
            session.rollback()
            self.logger.error(f"Database error: {str(e)}")
            return False
        finally:
            session.close()

    def insert_candidate(self, session, candidate_data):
        """Insert candidate data into candidates table and return candidate_id"""
        insert_stmt = self.candidates_table.insert().values(**candidate_data)
        result = session.execute(insert_stmt)
        session.commit()
        return result.inserted_primary_key[0]

    def insert_skills(self, session, candidate_id, skills):
        """Insert skills for a candidate"""
        for skill in skills:
            skill_data = {
                'candidate_id': candidate_id,
                'skill_name': skill,
                'proficiency_level': None  # Can be extended to extract proficiency
            }
            session.execute(self.skills_table.insert().values(**skill_data))
        session.commit()

    def insert_projects(self, session, candidate_id, projects):
        """Insert projects for a candidate"""
        for project in projects:
            project_data = {
                'candidate_id': candidate_id,
                'project_title': project.get('title'),
                'project_description': project.get('description'),
                'technologies_used': project.get('technologies'),
                'project_url': project.get('url')
            }
            session.execute(self.projects_table.insert().values(**project_data))
        session.commit()

    def insert_work_experience(self, session, candidate_id, work_experiences):
        """Insert work experience for a candidate"""
        for work in work_experiences:
            work_data = {
                'candidate_id': candidate_id,
                'company_name': work.get('company'),
                'job_title': work.get('title'),
                'start_date': work.get('start_date'),
                'end_date': work.get('end_date'),
                'description': work.get('description')
            }
            session.execute(self.work_experience_table.insert().values(**work_data))
        session.commit()

    def insert_certifications(self, session, candidate_id, certifications):
        """Insert certifications for a candidate"""
        for cert in certifications:
            cert_data = {
                'candidate_id': candidate_id,
                'certification_name': cert.get('name'),
                'issuing_organization': cert.get('organization'),
                'issue_date': cert.get('issue_date'),
                'expiration_date': cert.get('expiration_date'),
                'credential_url': cert.get('url')
            }
            session.execute(self.certifications_table.insert().values(**cert_data))
        session.commit()

    def insert_analysis_results(self, session, candidate_id, insights):
        """Insert analysis results for a candidate"""
        analysis_data = {
            'candidate_id': candidate_id,
            'analysis_date': datetime.datetime.now(),
            'insights': insights
        }
        session.execute(self.analysis_results_table.insert().values(**analysis_data))
        session.commit()

    def map_and_insert(self, file_name: str, analysis_data: dict):
        """Map extracted data to schema and insert into database"""
        if not hasattr(self, 'Session'):
            self.logger.warning("No database connection configured")
            return False

        session = self.Session()
        try:
            # Map candidate data
            candidate_data = {
                'name': None,
                'email': None,
                'phone': None,
                'linkedin': None,
                'github': None,
                'total_experience': None,
                'highest_qualification': None,
                'university': None,
                'location': None,
                'resume_text': analysis_data.get('raw_text')
            }

            # Extract basic contact info from analysis_data
            contacts = analysis_data.get('contacts', {})
            emails = contacts.get('emails', [])
            phones = contacts.get('phones', [])
            candidate_data['email'] = emails[0] if emails else None
            candidate_data['phone'] = phones[0] if phones else None

            # Extract name from entities PERSON if available
            entities = analysis_data.get('entities', {})
            persons = entities.get('PERSON', [])
            candidate_data['name'] = persons[0] if persons else None

            # TODO: Extract linkedin, github, total_experience, highest_qualification, university, location from text or entities

            # Insert candidate and get candidate_id
            candidate_id = self.insert_candidate(session, candidate_data)

            # Insert skills
            skills = analysis_data.get('skills', [])
            self.insert_skills(session, candidate_id, skills)

            # Insert projects - placeholder empty list for now
            projects = []  # Can be extracted from analysis_data if available
            self.insert_projects(session, candidate_id, projects)

            # Insert work experience - placeholder empty list for now
            work_experiences = []  # Can be extracted from analysis_data if available
            self.insert_work_experience(session, candidate_id, work_experiences)

            # Insert certifications - placeholder empty list for now
            certifications = []  # Can be extracted from analysis_data if available
            self.insert_certifications(session, candidate_id, certifications)

            # Insert analysis results
            insights = "Extracted data from resume"
            self.insert_analysis_results(session, candidate_id, insights)

            return True
        except Exception as e:
            session.rollback()
            self.logger.error(f"Error inserting mapped data: {str(e)}")
            return False
        finally:
            session.close()

    def process_files(self) -> Dict[str, str]:
        """Process all PDF/DOCX files, analyze with NLP and store in DB."""
        results = {}
        
        for filename in os.listdir(self.input_directory):
            file_path = os.path.join(self.input_directory, filename)
            
            if filename.lower().endswith('.pdf'):
                analysis = self.process_pdf(file_path)
            elif filename.lower().endswith('.docx'):
                analysis = self.process_docx(file_path)
            else:
                continue
                
            # Save to database
            if analysis:
                self.save_to_db(filename, analysis)
            
            results[filename] = {
                'analysis': analysis
            }
            
        return results

class EnhancedDocumentProcessor:
    def __init__(self):
        # Load NLP models
        self.nlp = spacy.load('en_core_web_lg')
        self.classifier = pipeline("text-classification", model="distilbert-base-uncased")
        
    def extract_structured_data(self, text):
        doc = self.nlp(text)
        
        # Enhanced section detection
        sections = {
            'experience': self._extract_experience(doc),
            'education': self._extract_education(doc),
            'skills': self._extract_skills(doc),
            'projects': self._extract_projects(doc)
        }
        return sections
        
    def _extract_skills(self, doc):
        # Use NER and pattern matching
        skills = []
        for ent in doc.ents:
            if ent.label_ == 'SKILL':
                skills.append(ent.text)
        return list(set(skills))

class JobMatcher:
    def __init__(self):
        self.vectorizer = TfidfVectorizer()
        
    def calculate_match_score(self, resume_text, job_description):
        # Convert texts to vectors
        vectors = self.vectorizer.fit_transform([resume_text, job_description])
        
        # Calculate cosine similarity
        similarity = cosine_similarity(vectors[0:1], vectors[1:2])[0][0]
        return similarity * 100  # Convert to percentage

def main():
    # Get the current script's directory
    current_dir = os.path.dirname(os.path.abspath(__file__))
    DATABASE_URL = "mysql+pymysql://root:123456@localhost:3306/resume_db"
    # Configure input directory relative to the script location
    input_dir = os.path.join(current_dir, '..', 'src', 'components', 'datafiles', 'Input_files', 'sample resume')
    
    # Create processor instance
    processor = DocumentProcessor(input_dir, None, db_uri=DATABASE_URL)
    
    # Process all files
    results = processor.process_files()
    
    # Print results
    print("\nProcessing Complete!")         
    print("-" * 50)
    for input_file, result in results.items():
        print(f"Processed {input_file}")

if __name__ == "__main__":
    main()