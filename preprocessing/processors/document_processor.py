import os
from typing import Dict, Optional
from PyPDF2 import PdfReader
from docx import Document
from sqlalchemy import create_engine, Table, MetaData
from sqlalchemy.orm import sessionmaker
import spacy
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import re

from preprocessing.dataextraction import Base, Resume
from preprocessing.models.patterns import (
    SECTION_PATTERNS, SKILL_KEYWORDS, 
    CONTACT_PATTERNS, DATE_PATTERNS
)
from preprocessing.utils.text_utils import clean_text, setup_logger

class DocumentProcessor:
    """A class to process PDF/DOCX files, extract text, analyze with NLP and store in DB."""
    
    def __init__(self, input_directory: str = None, output_directory: str = None, db_uri: str = None):
        self.logger = setup_logger(__name__)
        self.input_directory = input_directory
        self.output_directory = output_directory
        
        # Load spacy model once
        self.nlp_model = spacy.load('en_core_web_lg')
        
        # Initialize database connection
        if db_uri:
            self.engine = create_engine(db_uri)
            Base.metadata.create_all(self.engine)
            self.Session = sessionmaker(bind=self.engine)
        
        if output_directory and not os.path.exists(output_directory):
            os.makedirs(output_directory)

        # Define ORM models for the schema tables
        metadata = MetaData()
        if hasattr(self, 'engine'):
            metadata.reflect(bind=self.engine)

            self.candidates_table = metadata.tables.get('candidates')
            self.skills_table = metadata.tables.get('skills')
            self.projects_table = metadata.tables.get('projects')
            self.work_experience_table = metadata.tables.get('work_experience')
            self.certifications_table = metadata.tables.get('certifications')
            self.rankings_table = metadata.tables.get('rankings')
            self.analysis_results_table = metadata.tables.get('analysis_results')

    def process_pdf(self, file_path: str) -> Optional[dict]:
        """Enhanced PDF extraction with section detection"""
        try:
            texts = []
            with open(file_path, 'rb') as file:
                reader = PdfReader(file)
                for page in reader.pages:
                    texts.append(clean_text(page.extract_text()))
            text = "\n".join(texts)
            return self._analyze_resume(text)
        except Exception as e:
            self.logger.error(f"PDF processing error: {str(e)}")
            return None
            
    def process_docx(self, file_path: str) -> Optional[dict]:
        """Enhanced DOCX extraction with section detection"""
        try:
            doc = Document(file_path)
            text = "\n".join([clean_text(para.text) for para in doc.paragraphs])
            return self._analyze_resume(text)
        except Exception as e:
            self.logger.error(f"DOCX processing error: {str(e)}")
            return None

    def _analyze_with_nlp(self, text: str) -> dict:
        """Perform NLP analysis on resume text"""
        doc = self.nlp_model(text)
        
        # Extract entities
        entities = {
            'PERSON': [],
            'ORG': [],
            'DATE': [],
            'CGPA': [],
            'TECH': []
        }
        
        for ent in doc.ents:
            if ent.label_ in entities:
                entities[ent.label_].append(ent.text)
        
        # Extract skills using noun chunks and matching
        skills = []
        for chunk in doc.noun_chunks:
            if any(skill.lower() in chunk.text.lower() for skill in SKILL_KEYWORDS):
                skills.append(chunk.text)
        
        return {
            'entities': entities,
            'skills': list(set(skills))
        }

    def _analyze_resume(self, text: str) -> dict:
        """Enhanced resume analysis with NLP, contact info and dates"""
        # Detect sections
        sections = {
            name: bool(re.search(pattern, text, re.I))
            for name, pattern in SECTION_PATTERNS.items()
        }
        
        # Extract contact info
        contacts = {
            'emails': list(set(re.findall(CONTACT_PATTERNS['email'], text, re.I))),
            'phones': list(set(re.findall(CONTACT_PATTERNS['phone'], text))),
            'urls': list(set(re.findall(CONTACT_PATTERNS['url'], text, re.I)))
        }
        
        # Find dates in experience/education sections
        dates = {
            'education': re.findall(DATE_PATTERNS['year_range'], text, re.I),
            'experience': re.findall(DATE_PATTERNS['month_year'], text, re.I)
        }
        
        # Basic analysis
        analysis = {
            'raw_text': text,
            'sections': sections,
            'contacts': contacts,
            'dates': dates
        }
        
        # Add NLP analysis
        nlp_analysis = self._analyze_with_nlp(text)
        analysis.update(nlp_analysis)
        
        return analysis

    def save_to_db(self, file_name: str, analysis_data: dict) -> bool:
        """Save analysis results to database"""
        if not hasattr(self, 'Session'):
            self.logger.warning("No database connection configured")
            return False
            
        session = self.Session()
        try:
            resume = Resume(
                file_name=file_name,
                raw_text=analysis_data['raw_text'],
                extracted_data=analysis_data
            )
            session.add(resume)
            session.commit()
            return True
        except Exception as e:
            session.rollback()
            self.logger.error(f"Database error: {str(e)}")
            return False
        finally:
            session.close()

    def _process_single_file(self, file_path: str) -> Optional[dict]:
        """Process a single file based on its extension."""
        if file_path.lower().endswith('.pdf'):
            return self.process_pdf(file_path)
        elif file_path.lower().endswith('.docx'):
            return self.process_docx(file_path)
        else:
            self.logger.warning(f"Unsupported file type: {file_path}")
            return None

    def process_files(self) -> Dict[str, dict]:
        """Process all PDF/DOCX files, analyze with NLP and store in DB."""
        results = {}

        for filename in os.listdir(self.input_directory):
            file_path = os.path.join(self.input_directory, filename)
            analysis = self._process_single_file(file_path)

            if analysis:
                self.save_to_db(filename, analysis)

            results[filename] = {
                'analysis': analysis
            }

        self.logger.info(f"Processed {len(results)} files from {self.input_directory}")
        return results
